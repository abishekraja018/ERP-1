"""
Anna University CSE Department ERP System
Faculty (Staff) Views
"""

import json
from datetime import date, datetime
from django.utils import timezone

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Count

from .forms import (
    LeaveRequestForm, FeedbackForm, PublicationForm, 
    FacultyProfileEditForm, AccountUserForm, BulkAttendanceForm,
    QuestionPaperSubmissionForm
)
from .models import (
    Account_User, Faculty_Profile, Student_Profile, Course_Assignment,
    Attendance, LeaveRequest, Feedback, Notification, Publication,
    Announcement, AcademicYear, Semester, QuestionPaperAssignment,
    Timetable, TimetableEntry, TimeSlot, StructuredQuestionPaper, QPQuestion
)
from .utils.web_scrapper import fetch_acoe_updates
from .utils.cir_scrapper import fetch_cir_ticker_announcements


def check_faculty_permission(user, request=None):
    """
    Check if user is Faculty, Guest Faculty, or HOD in faculty mode.
    HOD is identified via Faculty_Profile.designation == 'HOD'.
    """
    if not user.is_authenticated:
        return False
    
    # Regular faculty or guest
    if user.role in ['FACULTY', 'GUEST']:
        return True
    
    # HOD (via Faculty_Profile.designation) can access faculty views
    if user.is_hod:
        return True
    
    return False


# =============================================================================
# DASHBOARD
# =============================================================================

@login_required
def staff_home(request):
    """Faculty Dashboard"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    
    # Get assigned courses
    assignments = Course_Assignment.objects.filter(faculty=faculty, is_active=True)
    total_courses = assignments.count()
    
    # Count students across all assigned batches
    total_students = 0
    for assignment in assignments:
        student_count = Student_Profile.objects.filter(batch_label=assignment.batch_label).count()
        total_students += student_count
    
    # Leave and attendance stats
    total_leaves = LeaveRequest.objects.filter(user=request.user).count()
    pending_leaves = LeaveRequest.objects.filter(user=request.user, status='PENDING').count()
    
    # Attendance stats per course
    course_list = []
    attendance_count_list = []
    for assignment in assignments:
        attendance_count = Attendance.objects.filter(assignment=assignment).count()
        course_list.append(assignment.course.title[:15])
        attendance_count_list.append(attendance_count)
    
    # Publications
    publications = Publication.objects.filter(faculty=faculty)
    total_publications = publications.count()
    unverified_publications = publications.filter(is_verified=False).count()
    
    # Notifications
    notifications = Notification.objects.filter(recipient=request.user, is_read=False)[:5]
    
    # Fetch announcements
    announcements = []
    try:
        acoe_updates = fetch_acoe_updates()
        announcements.extend(acoe_updates)
    except:
        pass
    
    try:
        cir_announcements = fetch_cir_ticker_announcements(limit=5)
        announcements.extend(cir_announcements)
    except:
        pass
    
    # Department announcements
    dept_announcements = Announcement.objects.filter(
        is_active=True, 
        audience__in=['ALL', 'FACULTY']
    )[:5]
    
    context = {
        'page_title': f'Faculty Dashboard - {faculty.user.full_name}',
        'faculty': faculty,
        'total_courses': total_courses,
        'total_students': total_students,
        'total_leaves': total_leaves,
        'pending_leaves': pending_leaves,
        'course_list': json.dumps(course_list),
        'attendance_count_list': json.dumps(attendance_count_list),
        'total_publications': total_publications,
        'unverified_publications': unverified_publications,
        'notifications': notifications,
        'announcements': announcements,
        'dept_announcements': dept_announcements,
        'assignments': assignments,
    }
    return render(request, 'staff_template/home_content.html', context)


# =============================================================================
# ATTENDANCE MANAGEMENT
# =============================================================================

@login_required
def staff_take_attendance(request):
    """Take attendance for assigned courses"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    assignments = Course_Assignment.objects.filter(faculty=faculty, is_active=True)
    
    context = {
        'assignments': assignments,
        'page_title': 'Take Attendance'
    }
    return render(request, 'staff_template/staff_take_attendance.html', context)


@csrf_exempt
@login_required
def get_students(request):
    """Get students for a course assignment"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    assignment_id = request.POST.get('assignment')
    try:
        assignment = get_object_or_404(Course_Assignment, id=assignment_id)
        
        # Get students matching the batch label
        students = Student_Profile.objects.filter(batch_label=assignment.batch_label).select_related('user')
        
        student_data = []
        for student in students:
            data = {
                "id": str(student.id),
                "name": student.user.full_name,
                "register_no": student.register_no
            }
            student_data.append(data)
        
        return JsonResponse(json.dumps(student_data), content_type='application/json', safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt
@login_required
def save_attendance(request):
    """Save attendance records"""
    if not check_faculty_permission(request.user):
        return HttpResponse("Permission denied", status=403)
    
    student_data = request.POST.get('student_ids')
    attendance_date = request.POST.get('date')
    assignment_id = request.POST.get('assignment')
    period = request.POST.get('period', 1)
    
    students = json.loads(student_data)
    
    try:
        assignment = get_object_or_404(Course_Assignment, id=assignment_id)
        
        for student_dict in students:
            student = get_object_or_404(Student_Profile, id=student_dict.get('id'))
            
            # Check if attendance already exists
            attendance, created = Attendance.objects.update_or_create(
                student=student,
                assignment=assignment,
                date=attendance_date,
                period=period,
                defaults={
                    'status': 'PRESENT' if student_dict.get('status') else 'ABSENT'
                }
            )
        
        return HttpResponse("OK")
    except Exception as e:
        return HttpResponse(f"Error: {str(e)}", status=400)


@login_required
def staff_update_attendance(request):
    """Update attendance records"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    assignments = Course_Assignment.objects.filter(faculty=faculty, is_active=True)
    
    context = {
        'assignments': assignments,
        'page_title': 'Update Attendance'
    }
    return render(request, 'staff_template/staff_update_attendance.html', context)


@csrf_exempt
@login_required
def get_student_attendance(request):
    """Get attendance data for editing"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    assignment_id = request.POST.get('assignment')
    attendance_date_id = request.POST.get('attendance_date_id')
    period = request.POST.get('period', 1)
    
    try:
        assignment = get_object_or_404(Course_Assignment, id=assignment_id)
        
        # Get the date from attendance record
        attendance_record = Attendance.objects.filter(id=attendance_date_id).first()
        if attendance_record:
            attendance_date = attendance_record.date
        else:
            return JsonResponse(json.dumps([]), content_type='application/json', safe=False)
        
        # Get students for this batch
        students = Student_Profile.objects.filter(batch_label=assignment.batch_label).select_related('user')
        
        student_data = []
        for student in students:
            # Check if attendance exists
            try:
                attendance = Attendance.objects.get(
                    student=student, 
                    assignment=assignment,
                    date=attendance_date,
                    period=period
                )
                status = attendance.status == 'PRESENT'
            except Attendance.DoesNotExist:
                status = False
            
            data = {
                "id": str(student.id),
                "name": student.user.full_name,
                "register_no": student.register_no,
                "status": status
            }
            student_data.append(data)
        
        return JsonResponse(json.dumps(student_data), content_type='application/json', safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt
@login_required
def update_attendance(request):
    """Update existing attendance records"""
    if not check_faculty_permission(request.user):
        return HttpResponse("Permission denied", status=403)
    
    student_data = request.POST.get('student_ids')
    attendance_date_id = request.POST.get('date')
    assignment_id = request.POST.get('assignment')
    period = request.POST.get('period', 1)
    
    students = json.loads(student_data)
    
    try:
        assignment = get_object_or_404(Course_Assignment, id=assignment_id)
        
        # Get the date from attendance record
        attendance_record = Attendance.objects.filter(id=attendance_date_id).first()
        if attendance_record:
            attendance_date = attendance_record.date
        else:
            return HttpResponse("Invalid attendance record", status=400)
        
        for student_dict in students:
            student = get_object_or_404(Student_Profile, id=student_dict.get('id'))
            
            Attendance.objects.update_or_create(
                student=student,
                assignment=assignment,
                date=attendance_date,
                period=period,
                defaults={
                    'status': 'PRESENT' if student_dict.get('status') else 'ABSENT'
                }
            )
        
        return HttpResponse("OK")
    except Exception as e:
        return HttpResponse(f"Error: {str(e)}", status=400)


# =============================================================================
# LEAVE MANAGEMENT
# =============================================================================

@login_required
def staff_apply_leave(request):
    """Apply for leave"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    form = LeaveRequestForm(request.POST or None, request.FILES or None)
    leave_history = LeaveRequest.objects.filter(user=request.user).order_by('-created_at')
    
    context = {
        'form': form,
        'leave_history': leave_history,
        'page_title': 'Apply for Leave'
    }
    
    if request.method == 'POST':
        if form.is_valid():
            try:
                leave = form.save(commit=False)
                leave.user = request.user
                leave.save()
                messages.success(request, "Leave application submitted for review")
                return redirect(reverse('staff_apply_leave'))
            except Exception as e:
                messages.error(request, f"Could not apply: {str(e)}")
        else:
            messages.error(request, "Please correct the errors in the form")
    
    return render(request, "staff_template/staff_apply_leave.html", context)


# =============================================================================
# FEEDBACK
# =============================================================================

@login_required
def staff_feedback(request):
    """Submit feedback"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    form = FeedbackForm(request.POST or None)
    feedbacks = Feedback.objects.filter(user=request.user).order_by('-created_at')
    
    context = {
        'form': form,
        'feedbacks': feedbacks,
        'page_title': 'Submit Feedback'
    }
    
    if request.method == 'POST':
        if form.is_valid():
            try:
                feedback = form.save(commit=False)
                feedback.user = request.user
                feedback.save()
                messages.success(request, "Feedback submitted successfully")
                return redirect(reverse('staff_feedback'))
            except Exception as e:
                messages.error(request, f"Could not submit: {str(e)}")
        else:
            messages.error(request, "Please correct the errors in the form")
    
    return render(request, "staff_template/staff_feedback.html", context)


# =============================================================================
# PROFILE
# =============================================================================

@login_required
def staff_view_profile(request):
    """View and update profile"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    user_form = AccountUserForm(request.POST or None, request.FILES or None, instance=request.user)
    profile_form = FacultyProfileEditForm(request.POST or None, instance=faculty)
    
    context = {
        'user_form': user_form,
        'profile_form': profile_form,
        'page_title': 'View/Update Profile'
    }
    
    if request.method == 'POST':
        if user_form.is_valid() and profile_form.is_valid():
            try:
                user = user_form.save(commit=False)
                password = user_form.cleaned_data.get('password')
                if password:
                    user.set_password(password)
                
                if 'profile_pic' in request.FILES:
                    fs = FileSystemStorage()
                    filename = fs.save(request.FILES['profile_pic'].name, request.FILES['profile_pic'])
                    user.profile_pic = fs.url(filename)
                
                user.save()
                profile_form.save()
                messages.success(request, "Profile Updated!")
                return redirect(reverse('staff_view_profile'))
            except Exception as e:
                messages.error(request, f"Error updating profile: {str(e)}")
        else:
            messages.error(request, "Invalid data provided")
    
    return render(request, "staff_template/staff_view_profile.html", context)


# =============================================================================
# NOTIFICATIONS
# =============================================================================

@csrf_exempt
@login_required
def staff_fcmtoken(request):
    """Update FCM token for push notifications"""
    token = request.POST.get('token')
    try:
        user = get_object_or_404(Account_User, id=request.user.id)
        user.fcm_token = token
        user.save()
        return HttpResponse("True")
    except Exception as e:
        return HttpResponse("False")


@login_required
def staff_view_notification(request):
    """View notifications"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    notifications = Notification.objects.filter(recipient=request.user).order_by('-created_at')
    
    # Mark as read
    notifications.filter(is_read=False).update(is_read=True)
    
    context = {
        'notifications': notifications,
        'page_title': "View Notifications"
    }
    return render(request, "staff_template/staff_view_notification.html", context)


# =============================================================================
# PUBLICATIONS
# =============================================================================

@login_required
def staff_add_publication(request):
    """Add new publication"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    form = PublicationForm(request.POST or None, request.FILES or None)
    publications = Publication.objects.filter(faculty=faculty).order_by('-created_at')
    
    context = {
        'form': form,
        'publications': publications,
        'page_title': 'Add Publication'
    }
    
    if request.method == 'POST':
        if form.is_valid():
            try:
                publication = form.save(commit=False)
                publication.faculty = faculty
                publication.save()
                messages.success(request, "Publication added successfully. Pending verification.")
                return redirect(reverse('staff_add_publication'))
            except Exception as e:
                messages.error(request, f"Could not add: {str(e)}")
        else:
            messages.error(request, "Please fill all required fields correctly")
    
    return render(request, "staff_template/staff_add_publication.html", context)


@login_required
def staff_view_publications(request):
    """View all publications"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    publications = Publication.objects.filter(faculty=faculty).order_by('-created_at')
    
    context = {
        'publications': publications,
        'page_title': 'My Publications'
    }
    return render(request, "staff_template/staff_view_publications.html", context)


# =============================================================================
# VIEW STUDENTS
# =============================================================================

@login_required
def staff_view_students(request):
    """View students in assigned courses"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    assignments = Course_Assignment.objects.filter(faculty=faculty, is_active=True)
    
    # Get unique batch labels
    batch_labels = assignments.values_list('batch_label', flat=True).distinct()
    
    # Get students in those batches
    students = Student_Profile.objects.filter(batch_label__in=batch_labels).select_related('user', 'advisor')
    
    context = {
        'students': students,
        'assignments': assignments,
        'page_title': 'View Students'
    }
    return render(request, "staff_template/staff_view_students.html", context)


# =============================================================================
# ATTENDANCE REPORTS
# =============================================================================

@login_required
def staff_view_attendance_report(request):
    """View attendance reports for assigned courses"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    assignments = Course_Assignment.objects.filter(faculty=faculty, is_active=True)
    
    selected_assignment = None
    attendance_data = []
    
    if request.GET.get('assignment'):
        selected_assignment = get_object_or_404(Course_Assignment, id=request.GET.get('assignment'))
        
        # Get students and their attendance
        students = Student_Profile.objects.filter(batch_label=selected_assignment.batch_label)
        
        for student in students:
            total = Attendance.objects.filter(student=student, assignment=selected_assignment).count()
            present = Attendance.objects.filter(
                student=student, 
                assignment=selected_assignment,
                status='PRESENT'
            ).count()
            
            percentage = (present / total * 100) if total > 0 else 0
            
            attendance_data.append({
                'student': student,
                'total': total,
                'present': present,
                'absent': total - present,
                'percentage': round(percentage, 2)
            })
    
    context = {
        'assignments': assignments,
        'selected_assignment': selected_assignment,
        'attendance_data': attendance_data,
        'page_title': 'Attendance Report'
    }
    return render(request, "staff_template/staff_attendance_report.html", context)


# =============================================================================
# QUESTION PAPER MANAGEMENT
# =============================================================================

@login_required
def staff_view_qp_assignments(request):
    """View question paper assignments for the faculty"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    
    # Get all QP assignments for this faculty
    assignments = QuestionPaperAssignment.objects.filter(
        assigned_faculty=faculty
    ).select_related(
        'course', 'academic_year', 'semester', 'regulation'
    ).order_by('-created_at')
    
    # Statistics
    stats = {
        'total': assignments.count(),
        'pending': assignments.filter(status__in=['ASSIGNED', 'IN_PROGRESS']).count(),
        'submitted': assignments.filter(status='SUBMITTED').count(),
        'approved': assignments.filter(status='APPROVED').count(),
        'rejected': assignments.filter(status='REJECTED').count(),
    }
    
    context = {
        'assignments': assignments,
        'stats': stats,
        'page_title': 'My Question Paper Assignments'
    }
    return render(request, "staff_template/staff_qp_assignments.html", context)


@login_required
def staff_submit_question_paper(request, qp_id):
    """Submit question paper for a specific assignment"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp_assignment = get_object_or_404(QuestionPaperAssignment, id=qp_id, assigned_faculty=faculty)
    
    # Check if already approved
    if qp_assignment.status == 'APPROVED':
        messages.info(request, "This question paper has already been approved.")
        return redirect('staff_view_qp_assignments')
    
    form = QuestionPaperSubmissionForm(request.POST or None, request.FILES or None, instance=qp_assignment)
    
    context = {
        'form': form,
        'qp_assignment': qp_assignment,
        'page_title': f'Submit QP - {qp_assignment.course.course_code}'
    }
    
    if request.method == 'POST':
        if form.is_valid():
            try:
                qp = form.save(commit=False)
                qp.status = 'SUBMITTED'
                qp.submitted_at = datetime.now()
                qp.save()
                
                # Create notification for HOD
                hod_users = Account_User.objects.filter(role='HOD', is_active=True)
                for hod in hod_users:
                    Notification.objects.create(
                        recipient=hod,
                        title='Question Paper Submitted',
                        message=f'{faculty.user.full_name} has submitted the {qp.get_exam_type_display()} question paper for {qp.course.course_code} - {qp.course.title}',
                        notification_type='INFO',
                        link=reverse('review_question_paper', args=[qp.id])
                    )
                
                messages.success(request, "Question paper submitted successfully! Waiting for review.")
                return redirect('staff_view_qp_assignments')
            except Exception as e:
                messages.error(request, f"Error submitting: {str(e)}")
        else:
            messages.error(request, "Please correct the errors in the form")
    
    return render(request, "staff_template/staff_submit_qp.html", context)


@login_required
def staff_view_qp_details(request, qp_id):
    """View details of a specific question paper assignment"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp_assignment = get_object_or_404(QuestionPaperAssignment, id=qp_id, assigned_faculty=faculty)
    
    context = {
        'qp_assignment': qp_assignment,
        'page_title': f'QP Details - {qp_assignment.course.course_code}'
    }
    return render(request, "staff_template/staff_qp_details.html", context)


# =============================================================================
# TIMETABLE VIEW
# =============================================================================

@login_required
def staff_view_timetable(request):
    """View faculty's teaching schedule across all batches"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    
    # Get all timetable entries for this faculty
    entries = TimetableEntry.objects.filter(
        faculty=faculty
    ).select_related(
        'timetable__academic_year', 'timetable__semester', 
        'course', 'time_slot'
    ).order_by('timetable', 'day', 'time_slot__slot_number')
    
    # Group entries by timetable (batch)
    timetable_entries = {}
    for entry in entries:
        tt_key = entry.timetable.id
        if tt_key not in timetable_entries:
            timetable_entries[tt_key] = {
                'timetable': entry.timetable,
                'entries': []
            }
        timetable_entries[tt_key]['entries'].append(entry)
    
    # Get time slots and days for consolidated view
    time_slots = TimeSlot.objects.all().order_by('slot_number')
    days = TimetableEntry.DAY_CHOICES
    
    # Create consolidated schedule (all batches combined)
    consolidated_schedule = {}
    for day_code, day_name in days:
        consolidated_schedule[day_code] = {}
        for slot in time_slots:
            slot_entries = entries.filter(day=day_code, time_slot=slot)
            consolidated_schedule[day_code][slot.slot_number] = list(slot_entries)
    
    context = {
        'faculty': faculty,
        'timetable_entries': timetable_entries,
        'consolidated_schedule': consolidated_schedule,
        'time_slots': time_slots,
        'days': days,
        'page_title': 'My Teaching Schedule'
    }
    return render(request, "staff_template/staff_timetable.html", context)


# =============================================================================
# STRUCTURED QUESTION PAPER (R2023 Format - Multi-field)
# =============================================================================

@login_required
def staff_list_structured_qps(request):
    """List all structured question papers and assignments for faculty"""
    from datetime import date
    
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    
    # Get faculty's QPs
    qps = StructuredQuestionPaper.objects.filter(faculty=faculty).select_related(
        'course', 'academic_year', 'semester', 'regulation'
    ).order_by('-created_at')
    
    # Get pending assignments (not yet completed)
    pending_assignments = QuestionPaperAssignment.objects.filter(
        assigned_faculty=faculty,
        status__in=['ASSIGNED', 'IN_PROGRESS']
    ).select_related('course', 'academic_year', 'semester', 'regulation').order_by('deadline')
    
    # Count stats for dashboard
    draft_count = qps.filter(status='DRAFT').count()
    submitted_count = qps.filter(status='SUBMITTED').count()
    approved_count = qps.filter(status='APPROVED').count()
    uploaded_count = qps.filter(is_uploaded=True).count()
    
    context = {
        'qps': qps,
        'pending_assignments': pending_assignments,
        'draft_count': draft_count,
        'submitted_count': submitted_count,
        'approved_count': approved_count,
        'uploaded_count': uploaded_count,
        'page_title': 'Question Paper Dashboard'
    }
    return render(request, "staff_template/list_structured_qps.html", context)


@login_required
def staff_upload_qp(request):
    """Upload a question paper document directly"""
    from main_app.forms import UploadQuestionPaperForm
    
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    
    if request.method == 'POST':
        form = UploadQuestionPaperForm(request.POST, request.FILES, faculty=faculty)
        if form.is_valid():
            try:
                qp = StructuredQuestionPaper(
                    faculty=faculty,
                    course=form.cleaned_data['course'],
                    academic_year=form.cleaned_data['academic_year'],
                    semester=form.cleaned_data['semester'],
                    regulation=form.cleaned_data['regulation'],
                    exam_month_year=form.cleaned_data['exam_month_year'],
                    uploaded_document=form.cleaned_data['uploaded_document'],
                    is_uploaded=True,
                    status='DRAFT'
                )
                qp.save()
                
                messages.success(request, f"Question paper uploaded successfully for {qp.course.course_code}!")
                return redirect('staff_list_structured_qps')
            except Exception as e:
                messages.error(request, f"Error uploading question paper: {str(e)}")
        else:
            messages.error(request, "Please correct the errors in the form.")
    else:
        form = UploadQuestionPaperForm(faculty=faculty)
    
    context = {
        'form': form,
        'page_title': 'Upload Question Paper'
    }
    return render(request, "staff_template/upload_qp.html", context)


@login_required
def staff_create_structured_qp(request, assignment_id=None):
    """Create structured question paper with formsets"""
    from main_app.forms import StructuredQuestionPaperForm, PartAFormSet, PartBFormSet, PartCFormSet
    
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    assignment = None
    
    if assignment_id:
        assignment = get_object_or_404(QuestionPaperAssignment, id=assignment_id, assigned_faculty=faculty)
        if hasattr(assignment, 'structured_qp'):
            messages.info(request, "Structured question paper already exists.")
            return redirect('staff_edit_structured_qp', qp_id=assignment.structured_qp.id)
    
    if request.method == 'POST':
        qp_form = StructuredQuestionPaperForm(request.POST)
        
        if qp_form.is_valid():
            qp = qp_form.save(commit=False)
            qp.faculty = faculty
            qp.qp_assignment = assignment
            qp.save()
            
            # Initialize formsets with the saved QP
            part_a_formset = PartAFormSet(request.POST, instance=qp, prefix='part_a')
            part_b_formset = PartBFormSet(request.POST, instance=qp, prefix='part_b')
            part_c_formset = PartCFormSet(request.POST, instance=qp, prefix='part_c')
            
            if part_a_formset.is_valid() and part_b_formset.is_valid() and part_c_formset.is_valid():
                # Save Part A questions
                for i, form in enumerate(part_a_formset):
                    if form.cleaned_data and form.cleaned_data.get('question_text'):
                        question = form.save(commit=False)
                        question.question_paper = qp
                        question.part = 'A'
                        question.question_number = i + 1
                        question.marks = 2
                        question.save()
                
                # Save Part B questions
                for i, form in enumerate(part_b_formset):
                    if form.cleaned_data and form.cleaned_data.get('question_text'):
                        question = form.save(commit=False)
                        question.question_paper = qp
                        question.part = 'B'
                        question.is_or_option = True
                        question.or_pair_number = 11 + (i // 2)
                        question.option_label = '(a)' if i % 2 == 0 else '(b)'
                        question.question_number = 11 + (i // 2)
                        question.marks = 13
                        question.save()
                
                # Save Part C question
                for form in part_c_formset:
                    if form.cleaned_data and form.cleaned_data.get('question_text'):
                        question = form.save(commit=False)
                        question.question_paper = qp
                        question.part = 'C'
                        question.question_number = 16
                        question.marks = 15
                        question.save()
                
                messages.success(request, "Question paper created successfully!")
                return redirect('staff_preview_structured_qp', qp_id=qp.id)
            else:
                messages.error(request, "Please check the questions for errors.")
        else:
            part_a_formset = PartAFormSet(request.POST, prefix='part_a')
            part_b_formset = PartBFormSet(request.POST, prefix='part_b')
            part_c_formset = PartCFormSet(request.POST, prefix='part_c')
    else:
        initial_data = {}
        if assignment:
            initial_data = {
                'course': assignment.course,
                'academic_year': assignment.academic_year,
                'semester': assignment.semester,
                'regulation': assignment.regulation,
            }
        
        qp_form = StructuredQuestionPaperForm(initial=initial_data)
        
        # Auto-populate formsets with empty forms
        part_a_initial = [{'question_number': i+1, 'marks': 2, 'part': 'A'} for i in range(10)]
        part_b_initial = [{'marks': 13, 'part': 'B'} for i in range(10)]
        part_c_initial = [{'question_number': 16, 'marks': 15, 'part': 'C'}]
        
        part_a_formset = PartAFormSet(prefix='part_a', queryset=QPQuestion.objects.none())
        part_b_formset = PartBFormSet(prefix='part_b', queryset=QPQuestion.objects.none())
        part_c_formset = PartCFormSet(prefix='part_c', queryset=QPQuestion.objects.none())
    
    context = {
        'qp_form': qp_form,
        'part_a_formset': part_a_formset,
        'part_b_formset': part_b_formset,
        'part_c_formset': part_c_formset,
        'assignment': assignment,
        'page_title': 'Create Structured Question Paper'
    }
    return render(request, "staff_template/create_structured_qp.html", context)


@login_required
def staff_edit_structured_qp(request, qp_id):
    """Edit existing structured question paper"""
    from main_app.forms import StructuredQuestionPaperForm, PartAFormSet, PartBFormSet, PartCFormSet
    
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    if qp.status not in ['DRAFT', 'REJECTED']:
        messages.warning(request, "Cannot edit submitted question paper.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)
    
    if request.method == 'POST':
        qp_form = StructuredQuestionPaperForm(request.POST, instance=qp)
        part_a_formset = PartAFormSet(request.POST, instance=qp, prefix='part_a')
        part_b_formset = PartBFormSet(request.POST, instance=qp, prefix='part_b')
        part_c_formset = PartCFormSet(request.POST, instance=qp, prefix='part_c')
        
        if qp_form.is_valid() and part_a_formset.is_valid() and part_b_formset.is_valid() and part_c_formset.is_valid():
            qp_form.save()
            
            # Update questions
            for i, form in enumerate(part_a_formset):
                if form.cleaned_data:
                    question = form.save(commit=False)
                    question.part = 'A'
                    question.question_number = i + 1
                    question.marks = 2
                    question.save()
            
            for i, form in enumerate(part_b_formset):
                if form.cleaned_data:
                    question = form.save(commit=False)
                    question.part = 'B'
                    question.is_or_option = True
                    question.or_pair_number = 11 + (i // 2)
                    question.option_label = '(a)' if i % 2 == 0 else '(b)'
                    question.question_number = 11 + (i // 2)
                    question.marks = 13
                    question.save()
            
            for form in part_c_formset:
                if form.cleaned_data:
                    question = form.save(commit=False)
                    question.part = 'C'
                    question.question_number = 16
                    question.marks = 15
                    question.save()
            
            messages.success(request, "Question paper updated successfully!")
            return redirect('staff_preview_structured_qp', qp_id=qp.id)
    else:
        qp_form = StructuredQuestionPaperForm(instance=qp)
        part_a_formset = PartAFormSet(instance=qp, prefix='part_a', queryset=qp.questions.filter(part='A').order_by('question_number'))
        part_b_formset = PartBFormSet(instance=qp, prefix='part_b', queryset=qp.questions.filter(part='B').order_by('or_pair_number', 'option_label'))
        part_c_formset = PartCFormSet(instance=qp, prefix='part_c', queryset=qp.questions.filter(part='C'))
    
    context = {
        'qp_form': qp_form,
        'part_a_formset': part_a_formset,
        'part_b_formset': part_b_formset,
        'part_c_formset': part_c_formset,
        'qp': qp,
        'page_title': 'Edit Structured Question Paper'
    }
    return render(request, "staff_template/create_structured_qp.html", context)


@login_required
def staff_preview_structured_qp(request, qp_id):
    """Preview structured question paper"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    # Get questions by part
    part_a_questions = qp.questions.filter(part='A').order_by('question_number')
    part_b_questions = qp.questions.filter(part='B').order_by('or_pair_number', 'option_label')
    part_c_questions = qp.questions.filter(part='C')
    
    # Group Part B into OR pairs
    part_b_pairs = []
    for i in range(11, 16):
        pair_questions = part_b_questions.filter(or_pair_number=i)
        if pair_questions.exists():
            part_b_pairs.append((i, pair_questions))
    
    # Calculate distribution
    distribution = {
        'co_distribution': {},
        'bloom_distribution': {},
        'l1_l2_total': 0,
        'l3_l4_total': 0,
        'l5_l6_total': 0,
        'l1_l2_percentage': 0,
        'l3_l4_percentage': 0,
        'l5_l6_percentage': 0,
    }
    
    all_questions = qp.questions.all()
    total_marks = sum(q.marks for q in all_questions)
    
    for co in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
        marks = sum(q.marks for q in all_questions if q.course_outcome == co)
        distribution['co_distribution'][co] = {
            'marks': marks,
            'percentage': (marks / total_marks * 100) if total_marks > 0 else 0
        }
    
    for bl in ['L1', 'L2', 'L3', 'L4', 'L5', 'L6']:
        marks = sum(q.marks for q in all_questions if q.bloom_level == bl)
        distribution['bloom_distribution'][bl] = {
            'marks': marks,
            'percentage': (marks / total_marks * 100) if total_marks > 0 else 0
        }
        
        if bl in ['L1', 'L2']:
            distribution['l1_l2_total'] += marks
        elif bl in ['L3', 'L4']:
            distribution['l3_l4_total'] += marks
        elif bl in ['L5', 'L6']:
            distribution['l5_l6_total'] += marks
    
    if total_marks > 0:
        distribution['l1_l2_percentage'] = distribution['l1_l2_total'] / total_marks * 100
        distribution['l3_l4_percentage'] = distribution['l3_l4_total'] / total_marks * 100
        distribution['l5_l6_percentage'] = distribution['l5_l6_total'] / total_marks * 100
    
    # Validation
    validation_errors = []
    if part_a_questions.count() != 10:
        validation_errors.append(f"Part A should have exactly 10 questions (found {part_a_questions.count()})")
    if part_b_questions.count() != 10:
        validation_errors.append(f"Part B should have exactly 10 questions (found {part_b_questions.count()})")
    if part_c_questions.count() != 1:
        validation_errors.append(f"Part C should have exactly 1 question (found {part_c_questions.count()})")
    
    if total_marks > 0:
        if distribution['l1_l2_percentage'] < 20 or distribution['l1_l2_percentage'] > 35:
            validation_errors.append(f"L1+L2 should be 20-35% (currently {distribution['l1_l2_percentage']:.1f}%)")
        if distribution['l3_l4_percentage'] < 40:
            validation_errors.append(f"L3+L4 should be ≥40% (currently {distribution['l3_l4_percentage']:.1f}%)")
        if distribution['l5_l6_percentage'] < 15 or distribution['l5_l6_percentage'] > 25:
            validation_errors.append(f"L5+L6 should be 15-25% (currently {distribution['l5_l6_percentage']:.1f}%)")
    
    can_submit = len(validation_errors) == 0 and qp.status in ['DRAFT', 'REJECTED']
    
    context = {
        'qp': qp,
        'part_a_questions': part_a_questions,
        'part_b_pairs': part_b_pairs,
        'part_c_questions': part_c_questions,
        'distribution': distribution,
        'validation_errors': validation_errors,
        'can_submit': can_submit,
        'page_title': 'Preview Question Paper'
    }
    return render(request, "staff_template/preview_structured_qp.html", context)


@login_required
def staff_submit_structured_qp(request, qp_id):
    """Submit question paper for HOD review"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    if request.method != 'POST':
        return redirect('staff_preview_structured_qp', qp_id=qp_id)
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    if qp.status not in ['DRAFT', 'REJECTED']:
        messages.warning(request, "Question paper already submitted.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)
    
    qp.status = 'SUBMITTED'
    qp.submitted_at = timezone.now()
    qp.save()
    
    # Update linked assignment status if exists
    if qp.qp_assignment:
        qp.qp_assignment.status = 'SUBMITTED'
        qp.qp_assignment.save()
    
    messages.success(request, "Question paper submitted successfully!")
    return redirect('staff_list_structured_qps')


@login_required
def staff_download_structured_qp(request, qp_id):
    """Download generated question paper"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    # Get questions by part
    part_a_questions = qp.questions.filter(part='A').order_by('question_number')
    part_b_questions = qp.questions.filter(part='B').order_by('or_pair_number', 'option_label')
    part_c_questions = qp.questions.filter(part='C')
    
    # Group Part B into OR pairs
    part_b_pairs = []
    for i in range(11, 16):
        pair_questions = list(part_b_questions.filter(or_pair_number=i))
        if pair_questions:
            part_b_pairs.append((i, pair_questions))
    
    # Generate document if not already generated
    if not qp.generated_document:
        from docx import Document
        from docx.shared import Inches, Pt, Twips, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsmap
        from docx.oxml import OxmlElement
        import io
        from django.core.files.base import ContentFile
        
        def set_cell_width(cell, width_inches):
            """Set exact cell width using XML"""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(width_inches * 1440)))  # Convert inches to twips
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        
        def set_table_fixed_layout(table, col_widths):
            """Set table to fixed layout with exact column widths"""
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Set table layout to fixed
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            
            # Set total table width
            total_width = sum(col_widths)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(int(total_width * 1440)))
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)
            
            # Add borders
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_el = OxmlElement(f'w:{border_name}')
                border_el.set(qn('w:val'), 'single')
                border_el.set(qn('w:sz'), '4')
                border_el.set(qn('w:space'), '0')
                border_el.set(qn('w:color'), '000000')
                tblBorders.append(border_el)
            tblPr.append(tblBorders)
            
            # Set column widths via tblGrid
            tblGrid = OxmlElement('w:tblGrid')
            for width in col_widths:
                gridCol = OxmlElement('w:gridCol')
                gridCol.set(qn('w:w'), str(int(width * 1440)))
                tblGrid.append(gridCol)
            tbl.insert(1, tblGrid)
            
            # Set cell widths for all rows
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    if idx < len(col_widths):
                        set_cell_width(cell, col_widths[idx])
        
        def set_cell_text(cell, text, bold=False, size=11, center=False):
            """Set cell text with formatting"""
            cell.text = text
            if cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                run.font.bold = bold
                run.font.size = Pt(size)
                run.font.name = 'Arial'
            if center:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        
        # Roll No - Right aligned with actual table boxes
        rollno_table = doc.add_table(rows=1, cols=11)
        rollno_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        
        # Set "Roll No." text in first cell (no border, wider width, bottom aligned)
        rollno_cell = rollno_table.rows[0].cells[0]
        rollno_cell.text = "Roll No."
        rollno_cell.paragraphs[0].runs[0].font.bold = True
        rollno_cell.paragraphs[0].runs[0].font.size = Pt(12)
        rollno_cell.paragraphs[0].runs[0].font.name = 'Arial'
        rollno_cell.width = Inches(0.9)
        # Set vertical alignment to bottom
        tc_rollno = rollno_cell._tc
        tcPr_rollno = tc_rollno.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'bottom')
        tcPr_rollno.append(vAlign)
        # Set fixed width for Roll No. cell
        tcW_rollno = OxmlElement('w:tcW')
        tcW_rollno.set(qn('w:w'), str(int(0.9 * 1440)))
        tcW_rollno.set(qn('w:type'), 'dxa')
        tcPr_rollno.append(tcW_rollno)
        
        # Add bordered boxes for each digit
        for i in range(1, 11):
            box_cell = rollno_table.rows[0].cells[i]
            box_cell.width = Inches(0.28)
            # Set cell borders
            tc = box_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border_el = OxmlElement(f'w:{border_name}')
                border_el.set(qn('w:val'), 'single')
                border_el.set(qn('w:sz'), '4')
                border_el.set(qn('w:space'), '0')
                border_el.set(qn('w:color'), '000000')
                tcBorders.append(border_el)
            tcPr.append(tcBorders)
            # Set fixed width
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(0.28 * 1440)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        
        doc.add_paragraph()  # spacing after roll no
        
        # University Header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.paragraph_format.space_before = Pt(0)
        header_para.paragraph_format.space_after = Pt(6)
        header_run = header_para.add_run("ANNA UNIVERSITY (UNIVERSITY DEPARTMENTS)")
        header_run.font.bold = True
        header_run.font.size = Pt(13)
        header_run.font.name = 'Arial'
        
        # Exam Info
        exam_para = doc.add_paragraph()
        exam_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        exam_para.paragraph_format.space_after = Pt(6)
        semester_type = getattr(qp.semester, 'semester_type', 'END SEMESTER')
        exam_run = exam_para.add_run(f"B.E. / B. Tech (Full Time) - {semester_type} EXAMINATIONS, {qp.exam_month_year.upper()}")
        exam_run.font.bold = True
        exam_run.font.size = Pt(12)
        exam_run.font.name = 'Arial'
        
        # Course Title
        course_title_para = doc.add_paragraph()
        course_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        course_title_para.paragraph_format.space_after = Pt(3)
        title_run = course_title_para.add_run(f"{qp.course.title.upper()}")
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        title_run.font.name = 'Arial'
        
        # Semester info
        sem_para = doc.add_paragraph()
        sem_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sem_para.paragraph_format.space_after = Pt(3)
        sem_num = getattr(qp.semester, 'semester_number', '')
        sem_run = sem_para.add_run(f"Semester {sem_num}" if sem_num else "Semester")
        sem_run.font.size = Pt(11)
        sem_run.font.name = 'Arial'
        
        # Course Code and Name
        code_para = doc.add_paragraph()
        code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        code_para.paragraph_format.space_after = Pt(3)
        code_run = code_para.add_run(f"{qp.course.course_code} - {qp.course.title}")
        code_run.font.bold = True
        code_run.font.size = Pt(12)
        code_run.font.name = 'Arial'
        
        # Regulation
        reg_para = doc.add_paragraph()
        reg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        reg_para.paragraph_format.space_after = Pt(12)
        reg_run = reg_para.add_run(f"(Regulation {qp.regulation.name})")
        reg_run.font.size = Pt(11)
        reg_run.font.name = 'Arial'
        
        # Time and Marks - use table for proper alignment
        time_marks_table = doc.add_table(rows=1, cols=2)
        time_marks_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        time_cell = time_marks_table.rows[0].cells[0]
        marks_cell = time_marks_table.rows[0].cells[1]
        
        time_cell.text = "Time: 3 hours"
        time_cell.paragraphs[0].runs[0].font.bold = True
        time_cell.paragraphs[0].runs[0].font.size = Pt(12)
        time_cell.paragraphs[0].runs[0].font.name = 'Arial'
        
        marks_cell.text = "Max. Marks: 100"
        marks_cell.paragraphs[0].runs[0].font.bold = True
        marks_cell.paragraphs[0].runs[0].font.size = Pt(12)
        marks_cell.paragraphs[0].runs[0].font.name = 'Arial'
        marks_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # spacing
        
        # CO Table - 2 columns with descriptions
        co_descriptions = [
            ('CO 1', getattr(qp, 'co1_description', '') or ''),
            ('CO 2', getattr(qp, 'co2_description', '') or ''),
            ('CO 3', getattr(qp, 'co3_description', '') or ''),
            ('CO 4', getattr(qp, 'co4_description', '') or ''),
            ('CO 5', getattr(qp, 'co5_description', '') or ''),
        ]
        
        co_table = doc.add_table(rows=5, cols=2)
        co_col_widths = [0.6, 6.5]  # inches
        set_table_fixed_layout(co_table, co_col_widths)
        
        for i, (label, desc) in enumerate(co_descriptions):
            set_cell_text(co_table.rows[i].cells[0], label, bold=True, size=11, center=True)
            set_cell_text(co_table.rows[i].cells[1], desc, bold=False, size=11, center=False)
        
        # BL Legend
        bl_para = doc.add_paragraph()
        bl_para.paragraph_format.space_before = Pt(6)
        bl_para.paragraph_format.space_after = Pt(12)
        
        bl_label = bl_para.add_run("BL – Bloom's Taxonomy Levels")
        bl_label.font.bold = True
        bl_label.font.size = Pt(11)
        bl_label.font.name = 'Arial'
        
        bl_para.add_run("\n")
        bl_desc = bl_para.add_run("(L1 - Remembering, L2 - Understanding, L3 - Applying, L4 - Analysing, L5 - Evaluating, L6 - Creating)")
        bl_desc.font.size = Pt(10)
        bl_desc.font.name = 'Arial'
        
        # ============ PART A ============
        part_a_title = doc.add_paragraph()
        part_a_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_a_title.paragraph_format.space_before = Pt(12)
        part_a_title.paragraph_format.space_after = Pt(3)
        a_run = part_a_title.add_run("PART- A (10 x 2 = 20 Marks)")
        a_run.font.bold = True
        a_run.font.size = Pt(12)
        a_run.font.name = 'Arial'
        a_run.font.underline = True
        
        part_a_subtitle = doc.add_paragraph()
        part_a_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_a_subtitle.paragraph_format.space_after = Pt(6)
        sub_run = part_a_subtitle.add_run("(Answer all Questions)")
        sub_run.font.size = Pt(11)
        sub_run.font.name = 'Arial'
        
        # Part A table - FIXED column widths
        part_a_col_widths = [0.55, 5.0, 0.55, 0.5, 0.5]  # Q.No, Questions, Marks, CO, BL
        part_a_table = doc.add_table(rows=11, cols=5)
        set_table_fixed_layout(part_a_table, part_a_col_widths)
        
        # Headers
        headers = ['Q. No', 'Questions', 'Marks', 'CO', 'BL']
        for i, header in enumerate(headers):
            set_cell_text(part_a_table.rows[0].cells[i], header, bold=True, size=11, center=True)
        
        # Questions
        for idx, q in enumerate(part_a_questions[:10], 1):
            row = part_a_table.rows[idx]
            set_cell_text(row.cells[0], str(q.question_number), size=11, center=True)
            set_cell_text(row.cells[1], q.question_text or '', size=11, center=False)
            set_cell_text(row.cells[2], str(q.marks), size=11, center=True)
            # Extract just the number from CO1, CO2, etc.
            co_num = (q.course_outcome or '').replace('CO', '').strip()
            set_cell_text(row.cells[3], co_num, size=11, center=True)
            set_cell_text(row.cells[4], q.bloom_level or '', size=11, center=True)
        
        # ============ PART B ============
        part_b_title = doc.add_paragraph()
        part_b_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_b_title.paragraph_format.space_before = Pt(12)
        part_b_title.paragraph_format.space_after = Pt(3)
        b_run = part_b_title.add_run("PART- B (5 x 13 = 65 Marks)")
        b_run.font.bold = True
        b_run.font.size = Pt(12)
        b_run.font.name = 'Arial'
        b_run.font.underline = True
        
        part_b_subtitle = doc.add_paragraph()
        part_b_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_b_subtitle.paragraph_format.space_after = Pt(6)
        sub_run = part_b_subtitle.add_run("(Answer any FIVE questions, choosing one from each unit)")
        sub_run.font.size = Pt(11)
        sub_run.font.name = 'Arial'
        
        # Part B table
        part_b_row_count = 1 + len(part_b_pairs) * 3  # header + (a, OR, b) for each pair
        part_b_table = doc.add_table(rows=1, cols=5)
        part_b_col_widths = [0.55, 5.0, 0.55, 0.5, 0.5]
        set_table_fixed_layout(part_b_table, part_b_col_widths)
        
        # Headers
        for i, header in enumerate(headers):
            set_cell_text(part_b_table.rows[0].cells[i], header, bold=True, size=11, center=True)
        
        # Questions
        for pair_num, questions in part_b_pairs:
            # Option (a)
            row_a = part_b_table.add_row()
            set_cell_width(row_a.cells[0], part_b_col_widths[0])
            set_cell_width(row_a.cells[1], part_b_col_widths[1])
            set_cell_width(row_a.cells[2], part_b_col_widths[2])
            set_cell_width(row_a.cells[3], part_b_col_widths[3])
            set_cell_width(row_a.cells[4], part_b_col_widths[4])
            
            set_cell_text(row_a.cells[0], f"{pair_num} (a)", size=11, center=True)
            q_a = next((q for q in questions if q.option_label == '(a)'), None)
            if q_a:
                set_cell_text(row_a.cells[1], q_a.question_text or '', size=11, center=False)
                co_num = (q_a.course_outcome or '').replace('CO', '').strip()
                set_cell_text(row_a.cells[3], co_num, size=11, center=True)
                set_cell_text(row_a.cells[4], q_a.bloom_level or '', size=11, center=True)
            set_cell_text(row_a.cells[2], '13', size=11, center=True)
            
            # OR row
            or_row = part_b_table.add_row()
            set_cell_width(or_row.cells[0], part_b_col_widths[0])
            set_cell_width(or_row.cells[1], part_b_col_widths[1])
            set_cell_width(or_row.cells[2], part_b_col_widths[2])
            set_cell_width(or_row.cells[3], part_b_col_widths[3])
            set_cell_width(or_row.cells[4], part_b_col_widths[4])
            set_cell_text(or_row.cells[0], '', size=11, center=True)
            set_cell_text(or_row.cells[1], 'OR', bold=True, size=11, center=True)
            set_cell_text(or_row.cells[2], '', size=11, center=True)
            set_cell_text(or_row.cells[3], '', size=11, center=True)
            set_cell_text(or_row.cells[4], '', size=11, center=True)
            
            # Option (b)
            row_b = part_b_table.add_row()
            set_cell_width(row_b.cells[0], part_b_col_widths[0])
            set_cell_width(row_b.cells[1], part_b_col_widths[1])
            set_cell_width(row_b.cells[2], part_b_col_widths[2])
            set_cell_width(row_b.cells[3], part_b_col_widths[3])
            set_cell_width(row_b.cells[4], part_b_col_widths[4])
            
            set_cell_text(row_b.cells[0], f"{pair_num} (b)", size=11, center=True)
            q_b = next((q for q in questions if q.option_label == '(b)'), None)
            if q_b:
                set_cell_text(row_b.cells[1], q_b.question_text or '', size=11, center=False)
                co_num = (q_b.course_outcome or '').replace('CO', '').strip()
                set_cell_text(row_b.cells[3], co_num, size=11, center=True)
                set_cell_text(row_b.cells[4], q_b.bloom_level or '', size=11, center=True)
            set_cell_text(row_b.cells[2], '13', size=11, center=True)
        
        # ============ PART C ============
        part_c_title = doc.add_paragraph()
        part_c_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_c_title.paragraph_format.space_before = Pt(12)
        part_c_title.paragraph_format.space_after = Pt(3)
        c_run = part_c_title.add_run("PART- C (1 x 15 = 15 Marks)")
        c_run.font.bold = True
        c_run.font.size = Pt(12)
        c_run.font.name = 'Arial'
        c_run.font.underline = True
        
        part_c_subtitle = doc.add_paragraph()
        part_c_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_c_subtitle.paragraph_format.space_after = Pt(6)
        sub_run = part_c_subtitle.add_run("(Q.No. 16 is compulsory)")
        sub_run.font.size = Pt(11)
        sub_run.font.name = 'Arial'
        
        # Part C table
        part_c_table = doc.add_table(rows=2, cols=5)
        part_c_col_widths = [0.55, 5.0, 0.55, 0.5, 0.5]
        set_table_fixed_layout(part_c_table, part_c_col_widths)
        
        # Headers
        for i, header in enumerate(headers):
            set_cell_text(part_c_table.rows[0].cells[i], header, bold=True, size=11, center=True)
        
        # Question
        if part_c_questions.exists():
            q = part_c_questions.first()
            row = part_c_table.rows[1]
            set_cell_text(row.cells[0], '16', size=11, center=True)
            set_cell_text(row.cells[1], q.question_text or '', size=11, center=False)
            set_cell_text(row.cells[2], str(q.marks), size=11, center=True)
            co_num = (q.course_outcome or '').replace('CO', '').strip()
            set_cell_text(row.cells[3], co_num, size=11, center=True)
            set_cell_text(row.cells[4], q.bloom_level or '', size=11, center=True)
        
        # Save to ContentFile
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        qp.generated_document.save(f'qp_{qp.id}.docx', ContentFile(doc_io.read()), save=True)
    
    from django.http import FileResponse
    return FileResponse(qp.generated_document.open('rb'), as_attachment=True, filename=qp.generated_document.name.split('/')[-1])


@login_required
def staff_manage_qp_answers(request, qp_id):
    """Page for managing answers for each question using AI generation"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    # Get questions by part
    part_a_questions = qp.questions.filter(part='A').order_by('question_number')
    part_b_questions = qp.questions.filter(part='B').order_by('question_number', 'option_label')
    part_c_questions = qp.questions.filter(part='C').order_by('question_number')
    
    # Group Part B questions by OR pairs
    part_b_pairs = []
    for pair_num in [11, 12, 13, 14, 15]:
        pair_questions = part_b_questions.filter(or_pair_number=pair_num).order_by('option_label')
        if pair_questions.exists():
            part_b_pairs.append((pair_num, list(pair_questions)))
    
    # Count answered questions
    all_questions = qp.questions.all()
    answered_count = all_questions.exclude(answer='').exclude(answer__isnull=True).count()
    total_questions = all_questions.count()
    
    context = {
        'page_title': f'Manage Answers - {qp.course.course_code}',
        'qp': qp,
        'part_a_questions': part_a_questions,
        'part_b_pairs': part_b_pairs,
        'part_c_questions': part_c_questions,
        'answered_count': answered_count,
        'total_questions': total_questions,
    }
    
    return render(request, 'staff_template/manage_qp_answers.html', context)


@login_required
@csrf_exempt
def staff_generate_answer_options(request):
    """AJAX endpoint to generate AI answer options for a question"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    try:
        data = json.loads(request.body)
        question_text = data.get('question', '').strip()
        marks = int(data.get('marks', 2))
        part_type = data.get('part', 'A')
        course_name = data.get('course_name', '')
        
        if not question_text:
            return JsonResponse({'error': 'Question text is required'}, status=400)
        
        from .utils.ai_answer_generator import generate_answer_options
        
        answers = generate_answer_options(
            question_text=question_text,
            marks=marks,
            part_type=part_type,
            course_name=course_name,
            num_options=4
        )
        
        return JsonResponse({'answers': answers})
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def staff_download_answer_key(request, qp_id):
    """Download answer key for a structured question paper"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    # Check if any answers exist
    questions_with_answers = qp.questions.exclude(answer='').exclude(answer__isnull=True)
    if not questions_with_answers.exists():
        messages.warning(request, "No answers have been added to this question paper yet.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)
    
    # Generate answer key document
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("ANSWER KEY")
    title_run.font.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Arial'
    
    # Course details
    details_para = doc.add_paragraph()
    details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details_run = details_para.add_run(f"{qp.course.course_code} - {qp.course.title}")
    details_run.font.bold = True
    details_run.font.size = Pt(12)
    details_run.font.name = 'Arial'
    
    exam_para = doc.add_paragraph()
    exam_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exam_run = exam_para.add_run(f"Examination: {qp.exam_month_year}")
    exam_run.font.size = Pt(11)
    exam_run.font.name = 'Arial'
    
    doc.add_paragraph()  # Spacing
    
    # Part A
    part_a_questions = qp.questions.filter(part='A').order_by('question_number')
    if part_a_questions.exists():
        part_a_title = doc.add_paragraph()
        part_a_run = part_a_title.add_run("PART A - Short Answers (2 marks each)")
        part_a_run.font.bold = True
        part_a_run.font.size = Pt(12)
        part_a_run.font.name = 'Arial'
        part_a_run.font.underline = True
        
        for q in part_a_questions:
            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{q.question_number}. {q.question_text}")
            q_run.font.bold = True
            q_run.font.size = Pt(11)
            q_run.font.name = 'Arial'
            
            # Answer
            if q.answer:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run(f"Answer: {q.answer}")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
            else:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run("Answer: [Not provided]")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
                a_run.font.italic = True
            
            doc.add_paragraph()  # Spacing
    
    # Part B
    part_b_questions = qp.questions.filter(part='B').order_by('or_pair_number', 'option_label')
    if part_b_questions.exists():
        part_b_title = doc.add_paragraph()
        part_b_run = part_b_title.add_run("PART B - Descriptive Answers (13 marks each)")
        part_b_run.font.bold = True
        part_b_run.font.size = Pt(12)
        part_b_run.font.name = 'Arial'
        part_b_run.font.underline = True
        
        for q in part_b_questions:
            # Question
            q_label = f"Q{q.or_pair_number}{q.option_label}" if q.option_label else f"Q{q.question_number}"
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"{q_label}. {q.question_text}")
            q_run.font.bold = True
            q_run.font.size = Pt(11)
            q_run.font.name = 'Arial'
            
            # Answer
            if q.answer:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run(f"Answer: {q.answer}")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
            else:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run("Answer: [Not provided]")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
                a_run.font.italic = True
            
            doc.add_paragraph()  # Spacing
    
    # Part C
    part_c_questions = qp.questions.filter(part='C')
    if part_c_questions.exists():
        part_c_title = doc.add_paragraph()
        part_c_run = part_c_title.add_run("PART C - Problem Solving (15 marks)")
        part_c_run.font.bold = True
        part_c_run.font.size = Pt(12)
        part_c_run.font.name = 'Arial'
        part_c_run.font.underline = True
        
        for q in part_c_questions:
            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q16. {q.question_text}")
            q_run.font.bold = True
            q_run.font.size = Pt(11)
            q_run.font.name = 'Arial'
            
            # Answer
            if q.answer:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run(f"Answer: {q.answer}")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
            else:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run("Answer: [Not provided]")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
                a_run.font.italic = True
    
    # Save to response
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    from django.http import FileResponse
    filename = f'Answer_Key_{qp.course.course_code}_{qp.exam_month_year.replace("/", "_")}.docx'
    return FileResponse(doc_io, as_attachment=True, filename=filename)


@login_required
@csrf_exempt
def staff_save_question_answer(request):
    """AJAX endpoint to save an answer for a question"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    try:
        data = json.loads(request.body)
        question_id = data.get('question_id')
        answer = data.get('answer', '').strip()
        
        if not question_id:
            return JsonResponse({'error': 'Question ID is required'}, status=400)
        
        faculty = get_object_or_404(Faculty_Profile, user=request.user)
        question = get_object_or_404(QPQuestion, id=question_id, question_paper__faculty=faculty)
        
        question.answer = answer
        question.save()
        
        return JsonResponse({'success': True, 'message': 'Answer saved successfully'})
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def staff_delete_qp_question(request, question_id):
    """Delete a question from a question paper (only for DRAFT status)"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    question = get_object_or_404(QPQuestion, id=question_id, question_paper__faculty=faculty)
    
    qp = question.question_paper
    
    # Only allow deletion for draft or rejected QPs
    if qp.status not in ['DRAFT', 'REJECTED']:
        messages.error(request, "Cannot delete questions from a submitted question paper.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)
    
    # Store info for message
    part = question.part
    q_num = question.question_number
    
    # Delete the question
    question.delete()
    
    messages.success(request, f"Question {q_num} from Part {part} deleted successfully.")
    
    # Redirect back to the referring page or preview
    referer = request.META.get('HTTP_REFERER')
    if referer:
        return redirect(referer)
    return redirect('staff_preview_structured_qp', qp_id=qp.id)


@login_required
@csrf_exempt
def staff_delete_qp_question_ajax(request):
    """AJAX endpoint to delete a question from a question paper"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    try:
        data = json.loads(request.body)
        question_id = data.get('question_id')
        
        if not question_id:
            return JsonResponse({'error': 'Question ID is required'}, status=400)
        
        faculty = get_object_or_404(Faculty_Profile, user=request.user)
        question = get_object_or_404(QPQuestion, id=question_id, question_paper__faculty=faculty)
        
        qp = question.question_paper
        
        # Only allow deletion for draft or rejected QPs
        if qp.status not in ['DRAFT', 'REJECTED']:
            return JsonResponse({'error': 'Cannot delete questions from a submitted question paper.'}, status=403)
        
        # Store info for response
        part = question.part
        q_num = question.question_number
        
        # Delete the question
        question.delete()
        
        return JsonResponse({
            'success': True, 
            'message': f'Question {q_num} from Part {part} deleted successfully.'
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def staff_delete_structured_qp(request, qp_id):
    """Delete an entire structured question paper (only for DRAFT status)"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    # Only allow deletion for draft or rejected QPs
    if qp.status not in ['DRAFT', 'REJECTED']:
        messages.error(request, "Cannot delete a submitted or approved question paper. Only draft or rejected question papers can be deleted.")
        return redirect('staff_list_structured_qps')
    
    if request.method == 'POST':
        # Store info for message
        course_code = qp.course.course_code
        exam_month_year = qp.exam_month_year
        
        # Delete all associated questions first, then the QP
        qp.questions.all().delete()
        qp.delete()
        
        messages.success(request, f"Question paper for {course_code} ({exam_month_year}) deleted successfully.")
        return redirect('staff_list_structured_qps')
    
    # GET request - show confirmation page
    context = {
        'qp': qp,
        'page_title': 'Delete Question Paper'
    }
    return render(request, "staff_template/delete_structured_qp_confirm.html", context)


@login_required
@csrf_exempt
def staff_delete_structured_qp_ajax(request):
    """AJAX endpoint to delete an entire structured question paper"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    try:
        data = json.loads(request.body)
        qp_id = data.get('qp_id')
        
        if not qp_id:
            return JsonResponse({'error': 'Question Paper ID is required'}, status=400)
        
        faculty = get_object_or_404(Faculty_Profile, user=request.user)
        qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
        
        # Only allow deletion for draft or rejected QPs
        if qp.status not in ['DRAFT', 'REJECTED']:
            return JsonResponse({'error': 'Cannot delete a submitted or approved question paper. Only draft or rejected question papers can be deleted.'}, status=403)
        
        # Store info for response
        course_code = qp.course.course_code
        exam_month_year = qp.exam_month_year
        
        # Delete all associated questions first, then the QP
        qp.questions.all().delete()
        qp.delete()
        
        return JsonResponse({
            'success': True, 
            'message': f'Question paper for {course_code} ({exam_month_year}) deleted successfully.'
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


